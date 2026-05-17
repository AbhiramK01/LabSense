#!/usr/bin/env python3
"""
Convert PROJECT_DOCUMENTATION.md to a formatted Word document
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from pygments import lex
    from pygments.lexers import get_lexer_by_name, TextLexer
    from pygments.token import Token
    HAS_PYGMENTS = True
except ImportError:
    HAS_PYGMENTS = False


CODE_COLORS = {
    'keyword': RGBColor(0x00, 0x00, 0xFF),
    'name': RGBColor(0x00, 0x00, 0x00),
    'string': RGBColor(0xA3, 0x55, 0x00),
    'comment': RGBColor(0x00, 0x80, 0x00),
    'number': RGBColor(0x80, 0x00, 0x80),
    'operator': RGBColor(0x00, 0x00, 0x00),
    'punctuation': RGBColor(0x00, 0x00, 0x00),
    'text': RGBColor(0x00, 0x00, 0x00),
}

def add_formatted_text(paragraph, text):
    """Add text with markdown formatting (bold, italic, code)"""
    # Split by markdown formatting
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    
    for part in parts:
        if not part:
            continue
        # Bold text
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        # Italic text
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        # Code text
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
        # Checkmark
        elif part == '✅':
            paragraph.add_run('✓')
        # Regular text
        else:
            paragraph.add_run(part)


def _token_color(token_type):
    if token_type in Token.Keyword:
        return CODE_COLORS['keyword']
    if token_type in Token.Name:
        return CODE_COLORS['name']
    if token_type in Token.Literal.String:
        return CODE_COLORS['string']
    if token_type in Token.Comment:
        return CODE_COLORS['comment']
    if token_type in Token.Literal.Number:
        return CODE_COLORS['number']
    if token_type in Token.Operator:
        return CODE_COLORS['operator']
    if token_type in Token.Punctuation:
        return CODE_COLORS['punctuation']
    return CODE_COLORS['text']


def add_code_block(doc, code_text, language=''):
    """Render a fenced code block with syntax colors."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)

    if language and HAS_PYGMENTS:
        try:
            lexer = get_lexer_by_name(language.lower())
        except Exception:
            lexer = TextLexer()

        for token_type, value in lex(code_text, lexer):
            run = p.add_run(value)
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = _token_color(token_type)
    else:
        for line in code_text.splitlines(True):
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

def parse_markdown_to_word(md_file, docx_file):
    """Convert markdown file to formatted Word document"""
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_list = False
    in_code_block = False
    code_language = ''
    code_lines = []

    def flush_code_block():
        nonlocal code_lines, code_language
        if code_lines:
            add_code_block(doc, ''.join(code_lines), code_language)
            code_lines = []
            code_language = ''
    
    while i < len(lines):
        line = lines[i].rstrip()

        # Fenced code block handling
        fence_match = re.match(r'^```\s*([a-zA-Z0-9_+-]*)\s*$', line)
        if fence_match:
            if in_code_block:
                flush_code_block()
                in_code_block = False
            else:
                in_code_block = True
                code_language = fence_match.group(1).strip()
            i += 1
            continue

        if in_code_block:
            code_lines.append(line + '\n')
            i += 1
            continue
        
        # Skip empty lines (but add spacing between sections)
        if not line:
            if in_list:
                in_list = False
            i += 1
            continue
        
        # Title (single #)
        if line.startswith('# ') and not line.startswith('##'):
            title = line[2:].strip()
            p = doc.add_heading(title, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            in_list = False
            i += 1
            continue
        
        # Level 1 heading (##)
        if line.startswith('## ') and not line.startswith('###'):
            heading = line[3:].strip()
            doc.add_heading(heading, level=1)
            in_list = False
            i += 1
            continue
        
        # Level 2 heading (###)
        if line.startswith('### '):
            heading = line[4:].strip()
            doc.add_heading(heading, level=2)
            in_list = False
            i += 1
            continue
        
        # Level 3 heading (####)
        if line.startswith('#### '):
            heading = line[5:].strip()
            doc.add_heading(heading, level=3)
            in_list = False
            i += 1
            continue
        
        # Horizontal rule (---)
        if line.startswith('---'):
            p = doc.add_paragraph()
            p.add_run('_' * 50).font.color.rgb = RGBColor(200, 200, 200)
            in_list = False
            i += 1
            continue
        
        # Bullet list item
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            in_list = True
            i += 1
            continue
        
        # Numbered list item
        match = re.match(r'^(\d+)\. ', line)
        if match:
            text = re.sub(r'^\d+\. ', '', line).strip()
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            in_list = True
            i += 1
            continue
        
        # Continuation of list (indented)
        if (line.startswith('   ') or line.startswith('\t')) and in_list:
            text = line.strip()
            p = doc.add_paragraph(style='List Bullet 2')
            add_formatted_text(p, text)
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, line)
        in_list = False
        i += 1
    
    # Save document
    flush_code_block()
    doc.save(docx_file)
    print(f"✅ Word document created: {docx_file}")

if __name__ == '__main__':
    import sys
    if len(sys.argv) >= 2:
        md_file = sys.argv[1]
        if len(sys.argv) >= 3:
            docx_file = sys.argv[2]
        else:
            docx_file = md_file.replace('.md', '.docx')
    else:
        md_file = 'PROJECT_DOCUMENTATION.md'
        docx_file = 'PROJECT_DOCUMENTATION.docx'
    
    try:
        parse_markdown_to_word(md_file, docx_file)
    except ImportError:
        print("❌ Error: python-docx library not found")
        print("📦 Please install: pip install python-docx")
        sys.exit(1)
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
