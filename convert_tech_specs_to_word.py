#!/usr/bin/env python3
"""
Convert TECHNICAL_SPECIFICATIONS.md to a formatted Word document
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_formatted_text(paragraph, text):
    """Add text with markdown formatting (bold, italic, code)"""
    # Split by markdown formatting
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    
    for part in parts:
        if not part:
            continue
        # Bold text
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        # Italic text
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        # Code text
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
        # Checkmark
        elif part == '✅':
            paragraph.add_run('✓')
        # Regular text
        else:
            paragraph.add_run(part)

def parse_markdown_to_word(md_file, docx_file):
    """Convert markdown file to formatted Word document"""
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_list = False
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines (but add spacing between sections)
        if not line:
            if in_list:
                in_list = False
            i += 1
            continue
        
        # Title (single #)
        if line.startswith('# ') and not line.startswith('##'):
            title = line[2:].strip()
            p = doc.add_heading(title, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            in_list = False
            i += 1
            continue
        
        # Level 1 heading (##)
        if line.startswith('## ') and not line.startswith('###'):
            heading = line[3:].strip()
            doc.add_heading(heading, level=1)
            in_list = False
            i += 1
            continue
        
        # Level 2 heading (###)
        if line.startswith('### '):
            heading = line[4:].strip()
            doc.add_heading(heading, level=2)
            in_list = False
            i += 1
            continue
        
        # Level 3 heading (####)
        if line.startswith('#### '):
            heading = line[5:].strip()
            doc.add_heading(heading, level=3)
            in_list = False
            i += 1
            continue
        
        # Horizontal rule (---)
        if line.startswith('---'):
            p = doc.add_paragraph()
            p.add_run('_' * 50).font.color.rgb = RGBColor(200, 200, 200)
            in_list = False
            i += 1
            continue
        
        # Bullet list item
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            in_list = True
            i += 1
            continue
        
        # Numbered list item
        match = re.match(r'^(\d+)\. ', line)
        if match:
            text = re.sub(r'^\d+\. ', '', line).strip()
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            in_list = True
            i += 1
            continue
        
        # Continuation of list (indented)
        if (line.startswith('   ') or line.startswith('\t')) and in_list:
            text = line.strip()
            p = doc.add_paragraph(style='List Bullet 2')
            add_formatted_text(p, text)
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, line)
        in_list = False
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✅ Word document created: {docx_file}")

if __name__ == '__main__':
    import sys
    md_file = 'TECHNICAL_SPECIFICATIONS.md'
    docx_file = 'TECHNICAL_SPECIFICATIONS.docx'
    
    try:
        parse_markdown_to_word(md_file, docx_file)
    except ImportError:
        print("❌ Error: python-docx library not found")
        print("📦 Please install: pip install python-docx")
        sys.exit(1)
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)







